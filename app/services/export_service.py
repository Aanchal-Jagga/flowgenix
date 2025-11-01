# # app/services/export_service.py
# import os
# import json
# from datetime import datetime
# from typing import Dict, Any
# from docx import Document
# from docx.shared import Pt


# # ========================
# # 🧠 MERGE OCR RESULTS
# # ========================
# def merge_ocr_results(text_data, formulas, symbols):
#     """
#     Combine OCR text (Azure), Math OCR formulas, and detected symbols
#     into a single structured result for LLM input.
#     """

#     # ✅ Handle Azure text safely
#     azure_text = ""
#     if isinstance(text_data, dict):
#         # Azure Vision response may contain 'analyzeResult', 'readResult', or 'text'
#         if "text" in text_data:
#             azure_text = text_data["text"]
#         elif "readResult" in text_data:
#             for block in text_data["readResult"].get("blocks", []):
#                 for line in block.get("lines", []):
#                     azure_text += line.get("text", "") + "\n"
#         elif "analyzeResult" in text_data and "readResults" in text_data["analyzeResult"]:
#             for page in text_data["analyzeResult"]["readResults"]:
#                 for line in page.get("lines", []):
#                     azure_text += line.get("text", "") + "\n"
#     elif isinstance(text_data, str):
#         azure_text = text_data

#     azure_text = azure_text.strip() if azure_text else "No text detected"

#     # ✅ Handle Math OCR formulas safely
#     if isinstance(formulas, list) and formulas:
#         math_section = "\n".join([f"Formula {i+1}: {f}" for i, f in enumerate(formulas)])
#     elif isinstance(formulas, str) and formulas.strip():
#         math_section = formulas
#     else:
#         math_section = "No formulas detected"

#     # ✅ Handle Symbols safely
#     symbols_section = []
#     if isinstance(symbols, dict) and "symbols" in symbols:
#         for s in symbols["symbols"]:
#             name = s.get("name", "Unknown")
#             conf = s.get("confidence", 0)
#             symbols_section.append(f"{name} (confidence: {conf:.2f})")

#     symbols_text = "\n".join(symbols_section) if symbols_section else "No symbols detected"

#     # ✅ Combine all structured info for LLM
#     combined_text = f"""
# 🧠 **Extracted Whiteboard Content**

# **Text (Azure OCR):**
# {azure_text}

# **Mathematical Formulas (Math OCR):**
# {math_section}

# **Symbols (Azure Vision Object Detection):**
# {symbols_text}
# """

#     return {
#         "merged_text": combined_text,
#         "summary": {
#             "text_detected": azure_text[:100] + ("..." if len(azure_text) > 100 else ""),
#             "num_formulas": len(formulas) if isinstance(formulas, list) else 0,
#             "num_symbols": len(symbols_section)
#         }
#     }


# # ========================
# # 🔶 CONVERT TO FLOWCHART
# # ========================
# def convert_to_flowchart_format(merged_data: Dict[str, Any]) -> Dict[str, Any]:
#     """
#     Convert structured OCR + math + symbol data into a flowchart-compatible JSON.
#     """
#     nodes = []
#     edges = []
#     node_id = 1
#     y_offset = -200
#     x_offset = 0

#     # Create nodes from text content
#     for item in merged_data.get("text", []):
#         nodes.append({
#             "id": node_id,
#             "label": item.get("content", "Untitled"),
#             "x": x_offset,
#             "y": y_offset
#         })
#         if node_id > 1:
#             edges.append({"from": node_id - 1, "to": node_id})
#         node_id += 1
#         y_offset += 150

#     # Add math formulas as additional nodes
#     if merged_data.get("formulas"):
#         nodes.append({
#             "id": node_id,
#             "label": f"Formula: {merged_data['formulas']}",
#             "x": x_offset,
#             "y": y_offset
#         })
#         edges.append({"from": node_id - 1, "to": node_id})
#         node_id += 1
#         y_offset += 150

#     # Add symbols as separate nodes
#     if merged_data.get("symbols"):
#         for symbol in merged_data["symbols"].get("detected_symbols", []):
#             nodes.append({
#                 "id": node_id,
#                 "label": f"Symbol: {symbol['type']}",
#                 "x": x_offset,
#                 "y": y_offset
#             })
#             edges.append({"from": node_id - 1, "to": node_id})
#             node_id += 1
#             y_offset += 150

#     # Add a final "End Process" node
#     nodes.append({
#         "id": node_id,
#         "label": "End Process",
#         "x": 0,
#         "y": y_offset
#     })
#     edges.append({"from": node_id - 1, "to": node_id})

#     return {"nodes": nodes, "edges": edges}


# # ========================
# # 🧾 EXPORT TO WORD (OLD)
# # ========================
# def export_to_word(structured_data: Dict[str, Any]) -> str:
#     """
#     Export structured notes to a Word document.
#     """
#     document = Document()

#     # Title
#     document.add_heading('FlowGenix Structured Notes', 0)

#     # Text Section
#     document.add_heading('Text Content', level=1)
#     for item in structured_data.get("text", []):
#         p = document.add_paragraph()
#         p.add_run(f"- {item.get('content', '')}").font.size = Pt(11)

#     # Formula Section
#     if structured_data.get("formulas"):
#         document.add_heading('Mathematical Formulas', level=1)
#         document.add_paragraph(structured_data["formulas"])

#     # Symbol Section
#     if structured_data.get("symbols"):
#         document.add_heading('Detected Symbols', level=1)
#         for symbol in structured_data["symbols"].get("detected_symbols", []):
#             document.add_paragraph(f"- {symbol['type']}")

#     # Save file
#     output_dir = "exports"
#     os.makedirs(output_dir, exist_ok=True)
#     filename = f"flowgenix_notes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
#     file_path = os.path.join(output_dir, filename)
#     document.save(file_path)

#     return file_path


# # ========================
# # 🧠 GENERATE NOTES DOCX (NEW, REVISED)
# # ========================
# def generate_notes_docx(
#     structured_notes: str,
#     explanations: Dict[str, str], # <-- NEW ARGUMENT
#     original_filename: str,
#     azure_data=None
# ):
#     """
#     Create a Word document (.docx) from:
#     - LLM structured notes (structured_notes)
#     - LLM deep explanations (explanations)
#     - Azure OCR extracted text (azure_data)
#     """
#     output_dir = "generated_notes"
#     os.makedirs(output_dir, exist_ok=True)

#     doc = Document()
#     style = doc.styles['Normal']
#     style.font.name = 'Calibri'
#     style.font.size = Pt(11)

#     doc.add_heading("FlowGenix: Whiteboard Analysis", level=0)
#     doc.add_paragraph(f"Original file: {original_filename}\n")

#     # --- Section 1: Structured Notes from LLM ---
#     doc.add_heading("Structured Whiteboard Notes", level=1)
#     if structured_notes:
#         doc.add_paragraph(structured_notes)
#     else:
#         doc.add_paragraph("No structured notes were generated by the LLM.")
#     doc.add_page_break()

#     # --- Section 2: Deep Topic Explanations (NEW) ---
#     doc.add_heading("Deep Topic Explanations", level=1)
#     if explanations:
#         for topic, explanation in explanations.items():
#             if "Error:" in explanation or "Info:" in explanation:
#                 continue # Skip errors
            
#             doc.add_heading(topic, level=2)
#             p = doc.add_paragraph(explanation)
#             p.style.font.size = Pt(10)
#             doc.add_paragraph() # Add a space
#     else:
#         doc.add_paragraph("No deep explanations were generated.")

#     # --- Section 3: Raw Azure OCR Text (for reference) ---
#     doc.add_heading("Raw Text (from Azure OCR)", level=1)
    
#     azure_text = ""
#     try:
#         if isinstance(azure_data, dict):
#             # Try to get the full text from the processed result
#             azure_text = azure_data.get("full_text", "")
#             if not azure_text:
#                 # Fallback to a simpler text key if full_text isn't present
#                 azure_text = azure_data.get("text", "No text field found.")
        
#         if not azure_text.strip():
#             azure_text = "No raw text detected from Azure OCR."
            
#     except Exception as e:
#         azure_text = f"[Error extracting Azure OCR text: {str(e)}]"

#     doc.add_paragraph(azure_text)

#     # --- Save the document ---
#     base_filename = os.path.splitext(original_filename)[0]
#     output_path = os.path.join(output_dir, f"{base_filename}_notes.docx")
    
#     try:
#         doc.save(output_path)
#         print(f"📄 Word document saved at: {output_path}")
#         return output_path
#     except Exception as e:
#         print(f"Error saving document: {e}")
#         # Try a fallback name
#         fallback_path = os.path.join(output_dir, f"fallback_{uuid.uuid4().hex}.docx")
#         doc.save(fallback_path)
#         return fallback_path
# app/services/export_service.py# app/services/export_service.py
import os
import json
import uuid
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ========================
# 🧠 MERGE OCR RESULTS
# ========================
def merge_ocr_results(text_data, formulas, symbols):
    """
    Combine OCR text (Azure), Math OCR formulas, and detected symbols
    into a single structured result for LLM input.
    """

    # ✅ Handle Azure text safely
    azure_text = ""
    if isinstance(text_data, dict):
        # Azure Vision response may contain 'analyzeResult', 'readResult', or 'text'
        if "text" in text_data:
            azure_text = text_data["text"]
        elif "readResult" in text_data:
            for block in text_data["readResult"].get("blocks", []):
                for line in block.get("lines", []):
                    azure_text += line.get("text", "") + "\n"
        elif "analyzeResult" in text_data and "readResults" in text_data["analyzeResult"]:
            for page in text_data["analyzeResult"]["readResults"]:
                for line in page.get("lines", []):
                    azure_text += line.get("text", "") + "\n"
    elif isinstance(text_data, str):
        azure_text = text_data

    azure_text = azure_text.strip() if azure_text else "No text detected"

    # ✅ Handle Math OCR formulas safely
    if isinstance(formulas, list) and formulas:
        math_section = "\n".join([f"Formula {i+1}: {f}" for i, f in enumerate(formulas)])
    elif isinstance(formulas, str) and formulas.strip():
        math_section = formulas
    else:
        math_section = "No formulas detected"

    # ✅ Handle Symbols safely
    symbols_section = []
    if isinstance(symbols, dict) and "symbols" in symbols:
        for s in symbols["symbols"]:
            name = s.get("name", "Unknown")
            conf = s.get("confidence", 0)
            symbols_section.append(f"{name} (confidence: {conf:.2f})")

    symbols_text = "\n".join(symbols_section) if symbols_section else "No symbols detected"

    # ✅ Combine all structured info for LLM
    combined_text = f"""
🧠 **Extracted Whiteboard Content**

**Text (Azure OCR):**
{azure_text}

**Mathematical Formulas (Math OCR):**
{math_section}

**Symbols (Azure Vision Object Detection):**
{symbols_text}
"""

    return {
        "merged_text": combined_text,
        "summary": {
            "text_detected": azure_text[:100] + ("..." if len(azure_text) > 100 else ""),
            "num_formulas": len(formulas) if isinstance(formulas, list) else 0,
            "num_symbols": len(symbols_section)
        }
    }


# ========================
# 🔶 CONVERT TO FLOWCHART
# ========================
def convert_to_flowchart_format(merged_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Convert structured OCR + math + symbol data into a flowchart-compatible JSON.
    """
    nodes = []
    edges = []
    node_id = 1
    y_offset = -200
    x_offset = 0

    # Create nodes from text content
    for item in merged_data.get("text", []):
        nodes.append({
            "id": node_id,
            "label": item.get("content", "Untitled"),
            "x": x_offset,
            "y": y_offset
        })
        if node_id > 1:
            edges.append({"from": node_id - 1, "to": node_id})
        node_id += 1
        y_offset += 150

    # Add math formulas as additional nodes
    if merged_data.get("formulas"):
        nodes.append({
            "id": node_id,
            "label": f"Formula: {merged_data['formulas']}",
            "x": x_offset,
            "y": y_offset
        })
        edges.append({"from": node_id - 1, "to": node_id})
        node_id += 1
        y_offset += 150

    # Add symbols as separate nodes
    if merged_data.get("symbols"):
        for symbol in merged_data["symbols"].get("detected_symbols", []):
            nodes.append({
                "id": node_id,
                "label": f"Symbol: {symbol['type']}",
                "x": x_offset,
                "y": y_offset
            })
            edges.append({"from": node_id - 1, "to": node_id})
            node_id += 1
            y_offset += 150

    # Add a final "End Process" node
    nodes.append({
        "id": node_id,
        "label": "End Process",
        "x": 0,
        "y": y_offset
    })
    edges.append({"from": node_id - 1, "to": node_id})

    return {"nodes": nodes, "edges": edges}


# ========================
# 🧾 EXPORT TO WORD (OLD)
# ========================
def export_to_word(structured_data: Dict[str, Any]) -> str:
    """
    Export structured notes to a Word document.
    """
    document = Document()

    # Title
    document.add_heading('FlowGenix Structured Notes', 0)

    # Text Section
    document.add_heading('Text Content', level=1)
    for item in structured_data.get("text", []):
        p = document.add_paragraph()
        p.add_run(f"- {item.get('content', '')}").font.size = Pt(11)

    # Formula Section
    if structured_data.get("formulas"):
        document.add_heading('Mathematical Formulas', level=1)
        document.add_paragraph(structured_data["formulas"])

    # Symbol Section
    if structured_data.get("symbols"):
        document.add_heading('Detected Symbols', level=1)
        for symbol in structured_data["symbols"].get("detected_symbols", []):
            document.add_paragraph(f"- {symbol['type']}")

    # Save file
    output_dir = "exports"
    os.makedirs(output_dir, exist_ok=True)
    filename = f"flowgenix_notes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(output_dir, filename)
    document.save(file_path)

    return file_path

# ========================
# ✨ NEW HELPER FUNCTION
# ========================
def add_markdown_to_doc(doc: Document, markdown_text: str):
    """
    Parses a Markdown string and adds it to the .docx document
    with proper formatting for headings, bullets, and code.
    """
    # Helper to add runs with backticks as code formatting
    def add_text_with_code(paragraph, text):
        parts = text.split('`')
        for i, part in enumerate(parts):
            if i % 2 == 1:  # This is a code block
                run = paragraph.add_run(part)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                paragraph.add_run(part)

    in_code_block = False
    code_block_text = ""

    lines = markdown_text.split('\n')
    for line in lines:
        stripped_line = line.strip()

        # Handle Code Blocks (```)
        if stripped_line.startswith('```'):
            if not in_code_block:
                in_code_block = True
            else:
                # End of code block, add it
                in_code_block = False
                p = doc.add_paragraph()
                p.style = 'Quote' # Use 'Quote' style for a simple indented block
                run = p.add_run(code_block_text.strip('\n'))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                code_block_text = ""
            continue
        
        if in_code_block:
            code_block_text += line + '\n'
            continue

        # Handle Headings
        if stripped_line.startswith('### '):
            p = doc.add_heading(level=3)
            add_text_with_code(p, stripped_line[4:])
        elif stripped_line.startswith('## '):
            p = doc.add_heading(level=2)
            add_text_with_code(p, stripped_line[3:])
        elif stripped_line.startswith('# '):
            p = doc.add_heading(level=1)
            add_text_with_code(p, stripped_line[2:])
        
        # Handle Bullet Points
        elif stripped_line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_text_with_code(p, stripped_line[2:])
        
        elif stripped_line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_text_with_code(p, stripped_line[2:])
        
        # Handle Numbered Lists
        elif len(stripped_line) > 2 and stripped_line[0].isdigit() and stripped_line[1] == '.':
            p = doc.add_paragraph(style='List Number')
            add_text_with_code(p, stripped_line[2:].strip())

        # Handle Horizontal Rules (---) --- THIS IS THE FIXED CODE ---
        elif stripped_line == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            # ❌ OLD, BUGGY CODE: pPr.insert_element_before(...)
            # ✅ NEW, SIMPLE CODE:
            pPr.append(pBdr) # Just append the border properties
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            # --- END OF FIX ---

        # Handle Empty Lines (for spacing)
        elif not stripped_line:
            doc.add_paragraph() # Add an empty paragraph

        # Handle Regular Text
        else:
            p = doc.add_paragraph()
            add_text_with_code(p, stripped_line)

# ========================
# 🧠 GENERATE NOTES DOCX (UPDATED)
# ========================
def generate_notes_docx(
    structured_notes: str,
    explanations: Dict[str, str],
    original_filename: str,
    azure_data: Dict[str, Any] = None  # Made optional
):
    """
    Create a Word document (.docx) from:
    - LLM structured notes (structured_notes)
    - LLM deep explanations (explanations)
    - Azure OCR extracted text (azure_data) - (Optional)
    """
    output_dir = "generated_notes"
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading("FlowGenix: Whiteboard Analysis", level=0)
    p = doc.add_paragraph()
    p.add_run(f"Original file: {original_filename}\n").italic = True


    # --- Section 1: Structured Notes from LLM ---
    doc.add_heading("Structured Whiteboard Notes", level=1)
    if structured_notes:
        add_markdown_to_doc(doc, structured_notes)
    else:
        doc.add_paragraph("No structured notes were generated by the LLM.")
    doc.add_page_break()

    # --- Section 2: Deep Topic Explanations ---
    doc.add_heading("Deep Topic Explanations", level=1)
    if explanations:
        for topic, explanation in explanations.items():
            if "Error:" in explanation or "Info:" in explanation:
                continue # Skip errors
            
            doc.add_heading(topic, level=2)
            # Call the parser for the explanation text as well
            add_markdown_to_doc(doc, explanation)
            doc.add_paragraph() # Add a space
    else:
        doc.add_paragraph("No deep explanations were generated.")

    # --- Section 3: Raw Azure OCR Text (if available) ---
    if azure_data: 
        doc.add_heading("Raw Text (from Azure OCR)", level=1)
        azure_text = ""
        try:
            if isinstance(azure_data, dict):
                azure_text = azure_data.get("full_text", "")
                if not azure_text:
                    azure_text = azure_data.get("text", "No text field found.")
            if not azure_text.strip():
                azure_text = "No raw text detected from Azure OCR."
        except Exception as e:
            azure_text = f"[Error extracting Azure OCR text: {str(e)}]"
        
        doc.add_paragraph(azure_text)

    # --- Save the document ---
    base_filename = os.path.splitext(original_filename)[0]
    output_path = os.path.join(output_dir, f"{base_filename}_notes.docx")
    
    try:
        doc.save(output_path)
        print(f"📄 Word document saved at: {output_path}")
        return output_path
    except Exception as e:
        print(f"Error saving document: {e}")
        # Try a fallback name
        fallback_path = os.path.join(output_dir, f"fallback_{uuid.uuid4().hex}.docx")
        doc.save(fallback_path)
        return fallback_path